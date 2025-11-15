import requests
import viktor as vkt

# AEC Data Model GraphQL endpoint
AEC_GRAPHQL_URL = "https://developer.api.autodesk.com/aec/graphql"


def execute_graphql(
    query: str, token: str, region: str, variables: dict = None, timeout: int = 30
):
    """
    Execute a GraphQL query against the Autodesk AEC Data Model API.

    Args:
        query: GraphQL query string
        token: OAuth2 access token
        region: Region identifier (e.g., 'US', 'EMEA')
        variables: Optional dictionary of GraphQL variables
        timeout: Request timeout in seconds

    Returns:
        Dictionary containing the response data
    """
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
        "Region": region,
    }
    payload = {"query": query, "variables": variables or {}}
    resp = requests.post(
        AEC_GRAPHQL_URL, headers=headers, json=payload, timeout=timeout
    )

    if resp.status_code != 200:
        raise RuntimeError(f"HTTP {resp.status_code}: {resp.text}")

    body = resp.json()
    if body.get("errors"):
        raise RuntimeError(f"GraphQL errors: {body['errors']}")

    return body.get("data", {})


class Parametrization(vkt.Parametrization):
    """Application input parameters."""

    file_field = vkt.FileField("Select file from ACC", flex=60)
    # AutodeskFileField to select a file from Autodesk Construction Cloud
    # autodesk_file = vkt.AutodeskFileField(
    #     "Select Autodesk File",
    #     oauth2_integration="autodesk-integration",
    #     description="Select a file from Autodesk Construction Cloud to view and analyze",
    # )

    # Dynamic array for selecting categories to check
    required_categories = vkt.DynamicArray(
        "Required Categories",
        description="Add categories that should be present in the model with custom colors",
        default=[
            {"category": "Structural Framing", "color": vkt.Color(255, 0, 0)},
            {"category": "Structural Columns", "color": vkt.Color(0, 0, 255)},
            {"category": "Walls", "color": vkt.Color(0, 255, 0)},
        ],
    )
    required_categories.category = vkt.OptionField(
        "Category",
        options=[
            "Structural Framing",
            "Structural Columns",
            "Structural Foundations",
            "Walls",
            "Floors",
            "Roofs",
            "Ceilings",
            "Doors",
            "Windows",
            "Stairs",
            "Railings",
            "Curtain Panels",
            "Curtain Wall Mullions",
            "Furniture",
            "Mechanical Equipment",
            "Plumbing Fixtures",
            "Lighting Fixtures",
            "Electrical Equipment",
            "Ducts",
            "Pipes",
        ],
    )
    required_categories.color = vkt.ColorField(
        "Highlight Color", default=vkt.Color(0, 255, 0)
    )

    # Download button for Category Summary report
    download_report = vkt.DownloadButton(
        "Download Category Report",
        method="download_category_report",
        description="Download a Word document with the Category Summary table",
        flex=60,
    )


class Controller(vkt.Controller):
    """Main application controller."""

    parametrization = Parametrization

    @vkt.AutodeskView("3D Model Viewer", duration_guess=5)
    def view_autodesk_model(self, params, **kwargs):
        """
        Display the selected Autodesk file in the 3D viewer.

        Args:
            params: User input parameters
            **kwargs: Additional arguments

        Returns:
            AutodeskResult containing the model to display
        """
        if not params.autodesk_file:
            raise vkt.UserError(
                "Please select an Autodesk file from the input field above"
            )

        # Initialize the OAuth2 integration and get access token
        integration = vkt.external.OAuth2Integration("autodesk-integration")
        token = integration.get_access_token()

        # Return the Autodesk viewer result
        return vkt.AutodeskResult(params.autodesk_file, access_token=token)

    @vkt.TableView("Family Instances", duration_guess=10)
    def view_family_instances(self, params, **kwargs):
        """
        Display a table of family instances from the model, including Structural Framing
        and other important Revit families.

        Args:
            params: User input parameters
            **kwargs: Additional arguments

        Returns:
            TableResult containing family instance data
        """
        if not params.autodesk_file:
            raise vkt.UserError(
                "Please select an Autodesk file from the input field above"
            )

        # Initialize the OAuth2 integration and get access token
        integration = vkt.external.OAuth2Integration("autodesk-integration")
        token = integration.get_access_token()

        # Get region and AEC Data Model element group ID from the Autodesk file
        region = params.autodesk_file.get_region(token)
        group_id = params.autodesk_file.get_aec_data_model_element_group_id(token)

        # Define categories to query (Structural Framing and other important families)
        categories = [
            "Structural Framing",
            "Structural Columns",
            "Walls",
            "Floors",
            "Doors",
            "Windows",
        ]

        # Collect all family instances
        all_instances = []

        for category in categories:
            vkt.progress_message(
                f"Fetching {category} instances...",
                percentage=categories.index(category) / len(categories) * 100,
            )

            # GraphQL query to get family instances for this category
            query = """
            query FamilyInstances($elementGroupId: ID!, $rsqlFilter: String!, $pagination: PaginationInput) {
              elementsByElementGroup(
                elementGroupId: $elementGroupId,
                filter: { query: $rsqlFilter },
                pagination: $pagination
              ) {
                pagination { cursor pageSize }
                results {
                  id
                  name
                  properties(filter: { names: ["Family Name", "Type Name"] }) {
                    results {
                      name
                      value
                    }
                  }
                }
              }
            }
            """

            # Construct RSQL filter for this category
            rsql_filter = f"property.name.category=='{category}' and 'property.name.Element Context'==Instance"

            # Fetch all instances with pagination
            cursor = None
            limit = 100

            while True:
                variables = {
                    "elementGroupId": group_id,
                    "rsqlFilter": rsql_filter,
                    "pagination": {"limit": limit}
                    if not cursor
                    else {"cursor": cursor, "limit": limit},
                }

                data = execute_graphql(query, token, region, variables)
                block = data.get("elementsByElementGroup", {}) or {}
                page_results = block.get("results", []) or []

                # Process each element in the page
                for element in page_results:
                    element_name = element.get("name", "Unknown")

                    # Extract Family Name and Type Name from properties
                    family_name = "Unknown"
                    type_name = "Unknown"

                    properties = element.get("properties", {}).get("results", []) or []
                    for prop in properties:
                        prop_name = prop.get("name", "")
                        prop_value = prop.get("value", "")

                        if prop_name == "Family Name":
                            family_name = prop_value if prop_value else "Unknown"
                        elif prop_name == "Type Name":
                            type_name = prop_value if prop_value else "Unknown"

                    # Add to results
                    all_instances.append(
                        {
                            "category": category,
                            "family_name": family_name,
                            "type_name": type_name,
                            "element_name": element_name,
                        }
                    )

                # Check pagination
                page = block.get("pagination", {}) or {}
                new_cursor = page.get("cursor")

                # Stop on empty cursor, repeated cursor, or empty page
                if not new_cursor or new_cursor == cursor or len(page_results) == 0:
                    break

                cursor = new_cursor

        # Prepare table data
        if not all_instances:
            # Return empty table if no instances found
            return vkt.TableResult(
                [["No family instances found in the selected categories"]],
                column_headers=["Message"],
            )

        # Create table rows
        table_data = []
        for instance in all_instances:
            table_data.append(
                [
                    instance["category"],
                    instance["family_name"],
                    instance["type_name"],
                    instance["element_name"],
                ]
            )

        # Define column headers
        column_headers = [
            vkt.TableHeader("Category", align="left"),
            vkt.TableHeader("Family Name", align="left"),
            vkt.TableHeader("Type Name", align="left"),
            vkt.TableHeader("Element Name", align="left"),
        ]

        return vkt.TableResult(
            table_data, column_headers=column_headers, enable_sorting_and_filtering=True
        )

    @vkt.TableView("Category Summary", duration_guess=10)
    def view_category_summary(self, params, **kwargs):
        """
        Display a summary table that shows the same categories as the dropdown list
        and cross-checks whether they are present in the model.

        Args:
            params: User input parameters
            **kwargs: Additional arguments

        Returns:
            TableResult showing which categories from the dropdown are present in the model
        """
        if not params.autodesk_file:
            raise vkt.UserError(
                "Please select an Autodesk file from the input field above"
            )

        # Initialize the OAuth2 integration and get access token
        integration = vkt.external.OAuth2Integration("autodesk-integration")
        token = integration.get_access_token()

        # Get region and AEC Data Model element group ID from the Autodesk file
        region = params.autodesk_file.get_region(token)
        group_id = params.autodesk_file.get_aec_data_model_element_group_id(token)

        # Extract required categories from dynamic array
        required_categories = set(row["category"] for row in params.required_categories)

        # Define the master list of categories (same as dropdown options)
        all_categories = [
            "Structural Framing",
            "Structural Columns",
            "Structural Foundations",
            "Walls",
            "Floors",
            "Roofs",
            "Ceilings",
            "Doors",
            "Windows",
            "Stairs",
            "Railings",
            "Curtain Panels",
            "Curtain Wall Mullions",
            "Furniture",
            "Mechanical Equipment",
            "Plumbing Fixtures",
            "Lighting Fixtures",
            "Electrical Equipment",
            "Ducts",
            "Pipes",
        ]

        vkt.progress_message("Fetching category counts from model...", percentage=10)

        # Query to get all distinct categories in the model with their counts
        query = """
        query UsedCategories($elementGroupId: ID!, $limit: Int!) {
          distinctPropertyValuesInElementGroupByName(
            elementGroupId: $elementGroupId
            name: "Category"
            filter: { query: "'property.name.Element Context'==Instance" }
          ) {
            results {
              values(limit: $limit) {
                value
                count
              }
            }
          }
        }
        """

        variables = {
            "elementGroupId": group_id,
            "limit": 1000,  # High limit to get all categories
        }

        try:
            data = execute_graphql(query, token, region, variables)
            block = data.get("distinctPropertyValuesInElementGroupByName") or {}
            results_list = block.get("results") or []

            # Create a dictionary of category counts from the model
            model_category_counts = {}
            for r in results_list:
                values = r.get("values") or []
                for v in values:
                    category_name = v.get("value", "")
                    element_count = v.get("count", 0)
                    if category_name:
                        model_category_counts[category_name] = element_count

        except Exception as e:
            raise vkt.UserError(f"Failed to fetch categories from model: {str(e)}")

        vkt.progress_message("Preparing category summary...", percentage=80)

        # Prepare table data with visual indicators
        table_data = []
        for category_name in all_categories:
            # Check if category is in the model
            element_count = model_category_counts.get(category_name, 0)
            in_model = element_count > 0

            # Check if category is in required categories
            in_contract = category_name in required_categories

            # Determine status symbol and description
            if in_contract and in_model:
                status_symbol = "✓"
                status_text = "Present in contract and model"
                status_color = vkt.Color(0, 128, 0)  # Green
            elif in_contract and not in_model:
                status_symbol = "✗"
                status_text = "In contract but not in model"
                status_color = vkt.Color(255, 165, 0)  # Orange
            elif not in_contract and in_model:
                status_symbol = "✗"
                status_text = "Missing in the contract"
                status_color = vkt.Color(255, 0, 0)  # Red
            else:  # not in_contract and not in_model
                status_symbol = "✗"
                status_text = "Not in contract, not in model"
                status_color = vkt.Color(128, 128, 128)  # Gray

            # Create colored cells for better visualization
            status_cell = vkt.TableCell(
                status_symbol, text_color=status_color, text_style="bold"
            )

            table_data.append([category_name, status_cell, element_count, status_text])

        # Define column headers
        column_headers = [
            vkt.TableHeader("Category", align="left"),
            vkt.TableHeader("Status", align="center"),
            vkt.TableHeader("Element Count", align="right"),
            vkt.TableHeader("Description", align="left"),
        ]

        return vkt.TableResult(
            table_data, column_headers=column_headers, enable_sorting_and_filtering=True
        )

    @vkt.WebView("Colored Category View", duration_guess=15)
    def view_colored_categories(self, params, **kwargs):
        """
        Display the Autodesk model with categories highlighted in custom colors
        based on the dynamic field selections.
        """
        if not params.autodesk_file:
            raise vkt.UserError(
                "Please select an Autodesk file from the input field above"
            )

        # Initialize the OAuth2 integration and get access token
        integration = vkt.external.OAuth2Integration("autodesk-integration")
        token = integration.get_access_token()

        # Get the URN from the Autodesk file and encode it properly
        autodesk_file = params.autodesk_file
        region = autodesk_file.get_region(token)
        group_id = autodesk_file.get_aec_data_model_element_group_id(token)

        # Get the latest version URN and encode it like in your working example
        latest_version = autodesk_file.get_latest_version(token)
        urn = latest_version.urn
        import base64

        urn_bs64 = base64.urlsafe_b64encode(urn.encode()).decode().rstrip("=")

        vkt.progress_message(
            "Fetching element external IDs for selected categories...", percentage=20
        )

        # Build a list of external IDs with their colors for each category
        external_ids_with_colors = []

        for row in params.required_categories:
            category_name = row["category"]
            color = row["color"]

            # Convert VIKTOR Color to hex format
            color_hex = color.hex

            vkt.progress_message(f"Fetching {category_name} elements...", percentage=30)

            # GraphQL query to get element IDs and their external IDs
            query = """
            query CategoryElements($elementGroupId: ID!, $rsqlFilter: String!, $pagination: PaginationInput) {
            elementsByElementGroup(
                elementGroupId: $elementGroupId,
                filter: { query: $rsqlFilter },
                pagination: $pagination
            ) {
                pagination { cursor pageSize }
                results {
                id
                name
                alternativeIdentifiers {
                    externalElementId
                }
                }
            }
            }
            """

            # Construct RSQL filter for this category
            rsql_filter = f"property.name.category=='{category_name}' and 'property.name.Element Context'==Instance"

            # Fetch all elements with pagination
            cursor = None
            limit = 100

            while True:
                variables = {
                    "elementGroupId": group_id,
                    "rsqlFilter": rsql_filter,
                    "pagination": {"limit": limit}
                    if not cursor
                    else {"cursor": cursor, "limit": limit},
                }

                try:
                    data = execute_graphql(query, token, region, variables)
                    block = data.get("elementsByElementGroup", {}) or {}
                    page_results = block.get("results", []) or []

                    # Collect external IDs with their colors
                    for element in page_results:
                        # Get External ID from alternativeIdentifiers
                        alt_ids = element.get("alternativeIdentifiers", {})
                        external_id = alt_ids.get("externalElementId")

                        # If External ID found, add it with its color
                        if external_id:
                            # Create a single-key object as expected by the viewer script
                            external_ids_with_colors.append({external_id: color_hex})

                    # Check pagination
                    page = block.get("pagination", {}) or {}
                    new_cursor = page.get("cursor")

                    # Stop on empty cursor, repeated cursor, or empty page
                    if not new_cursor or new_cursor == cursor or len(page_results) == 0:
                        break

                    cursor = new_cursor

                except Exception as e:
                    vkt.UserMessage.warning(
                        f"Could not fetch elements for category '{category_name}': {str(e)}"
                    )
                    break

        vkt.progress_message("Preparing viewer...", percentage=80)

        # Convert Python list to JSON string for JavaScript
        import json

        external_ids_json = json.dumps(external_ids_with_colors)

        # Use the same HTML template approach as your working example
        html_template = """<!DOCTYPE html>
    <html>
    <head>
    <meta charset="utf-8" />
    <title>APS Viewer - Colored Categories</title>
    <link rel="stylesheet" href="https://developer.api.autodesk.com/modelderivative/v2/viewers/7.*/style.min.css" type="text/css">
    <script src="https://developer.api.autodesk.com/modelderivative/v2/viewers/7.*/viewer3D.min.js"></script>
    <style>
        html, body, #apsViewerDiv { width: 100%; height: 100%; margin: 0; }
    </style>
    </head>
    <body>
    <div id="apsViewerDiv"></div>

    <script>
    // Injected external IDs from backend
    var EXTERNAL_IDS = EXTERNAL_IDS_PLACEHOLDER;
    // Replace these three in your backend
    var ACCESS_TOKEN = 'APS_TOKEN_PLACEHOLDER';
    var DOCUMENT_URN = 'urn:URN_PLACEHOLDER';
    
    console.log('External IDs:', EXTERNAL_IDS);

    // Disable analytics for sandbox iframes
    try { Autodesk.Viewing.Private.analytics.optOut(); } catch (e) {}

    let viewer = null;
    let modelLoaded = null;

    Autodesk.Viewing.Initializer(
        { env: 'AutodeskProduction2', api: 'streamingV2', accessToken: ACCESS_TOKEN },
        function () {
        const container = document.getElementById('apsViewerDiv');
        viewer = new Autodesk.Viewing.GuiViewer3D(container, { disableBimWalkInfoIcon: true });
        viewer.start();
        console.log('Viewer started');

        if (!DOCUMENT_URN) {
            console.error('Missing URN');
            return;
        }

        Autodesk.Viewing.Document.load(
            DOCUMENT_URN,
            function onSuccess(doc) {
            const node = doc.getRoot().getDefaultGeometry();
            if (!node) { console.warn('No default geometry'); return; }
            viewer.loadDocumentNode(doc, node, { keepCurrentModels: false }).then(function (model) {
                modelLoaded = model;
                console.log('Model loaded');

                // Apply filtering and coloring
                modelLoaded.getExternalIdMapping(
                function onMap(map) {
                    if (!map) {
                    console.warn('externalId map not available');
                    return;
                    }
                    
                    const dbIds = [];
                    const extColorMap = buildExternalColorMap(EXTERNAL_IDS);
                    const rev = {}; // dbId -> externalId
                    
                    for (const extId in map) {
                    if (Object.prototype.hasOwnProperty.call(map, extId)) {
                        rev[ map[extId] ] = extId;
                    }
                    }
                    
                    // Collect dbIds for all external IDs in EXTERNAL_IDS
                    for (const extId in extColorMap) {
                    if (map[extId]) {
                        dbIds.push(map[extId]);
                    }
                    }
                    
                    if (dbIds.length === 0) {
                    console.info('No dbIds match EXTERNAL_IDS');
                    return;
                    }
                    
                    viewer.clearThemingColors();
                    viewer.isolate(dbIds);
                    viewer.fitToView(dbIds);
                    
                    // Apply colors
                    const THREE_REF = (Autodesk && Autodesk.Viewing && Autodesk.Viewing.Private && Autodesk.Viewing.Private.THREE) || window.THREE;
                    const defaultV4 = colorStringToVec4('green', 0.85, THREE_REF);
                    
                    for (let i = 0; i < dbIds.length; i++) {
                    viewer.setThemingColor(dbIds[i], defaultV4, modelLoaded, false);
                    }
                    
                    // Apply overrides from extColorMap
                    for (let i = 0; i < dbIds.length; i++) {
                    const dbId = dbIds[i];
                    const ext = rev[dbId];
                    if (ext && extColorMap[ext]) {
                        const v4 = colorStringToVec4(extColorMap[ext], 0.95, THREE_REF);
                        viewer.setThemingColor(dbId, v4, modelLoaded, false);
                    }
                    }
                    viewer.impl.invalidate(true, true, true);
                },
                function onErr(err) {
                    console.error('getExternalIdMapping failed', err);
                }
                );
            });
            },
            function onFailure(code, message) {
            console.error('Document load failed:', code, message);
            }
        );
        }
    );

    // Helper functions
    function buildExternalColorMap(list) {
        const map = Object.create(null);
        if (!Array.isArray(list)) return map;
        for (let i = 0; i < list.length; i++) {
        const obj = list[i];
        if (obj && typeof obj === 'object') {
            const keys = Object.keys(obj);
            if (keys.length === 1) {
            const extId = String(keys[0]);
            const color = String(obj[extId] || '').trim();
            if (extId && color) map[extId] = color;
            }
        }
        }
        return map;
    }

    function colorStringToVec4(str, alphaDefault, THREE_REF) {
        let a = typeof alphaDefault === 'number' ? alphaDefault : 0.85;
        if (!str) return new THREE_REF.Vector4(0, 1, 0, a);

        let s = String(str).trim().toLowerCase();
        
        // #rrggbb
        if (s.startsWith('#') && s.length === 7) {
        const r = parseInt(s.slice(1, 3), 16) / 255;
        const g = parseInt(s.slice(3, 5), 16) / 255;
        const b = parseInt(s.slice(5, 7), 16) / 255;
        return new THREE_REF.Vector4(r, g, b, a);
        }

        // fallback to green
        return new THREE_REF.Vector4(0, 1, 0, a);
    }
    </script>
    </body>
    </html>"""

        # Use placeholder replacement like in your working example
        html = html_template.replace("APS_TOKEN_PLACEHOLDER", token)
        html = html.replace("URN_PLACEHOLDER", urn_bs64)
        html = html.replace("EXTERNAL_IDS_PLACEHOLDER", external_ids_json)

        return vkt.WebResult(html=html)

    @vkt.DataView("Category Data Summary", duration_guess=10)
    def view_category_data(self, params, **kwargs):
        """
        Display a data summary showing which categories from the dropdown are present in the model.

        Args:
            params: User input parameters
            **kwargs: Additional arguments

        Returns:
            DataResult with category status information
        """
        if not params.autodesk_file:
            raise vkt.UserError(
                "Please select an Autodesk file from the input field above"
            )

        # Initialize the OAuth2 integration and get access token
        integration = vkt.external.OAuth2Integration("autodesk-integration")
        token = integration.get_access_token()

        # Get region and AEC Data Model element group ID from the Autodesk file
        region = params.autodesk_file.get_region(token)
        group_id = params.autodesk_file.get_aec_data_model_element_group_id(token)

        # Extract required categories from dynamic array
        required_categories = set(row["category"] for row in params.required_categories)

        # Define the master list of categories (same as dropdown options)
        all_categories = [
            "Structural Framing",
            "Structural Columns",
            "Structural Foundations",
            "Walls",
            "Floors",
            "Roofs",
            "Ceilings",
            "Doors",
            "Windows",
            "Stairs",
            "Railings",
            "Curtain Panels",
            "Curtain Wall Mullions",
            "Furniture",
            "Mechanical Equipment",
            "Plumbing Fixtures",
            "Lighting Fixtures",
            "Electrical Equipment",
            "Ducts",
            "Pipes",
        ]

        vkt.progress_message("Fetching category counts from model...", percentage=10)

        # Query to get all distinct categories in the model with their counts
        query = """
        query UsedCategories($elementGroupId: ID!, $limit: Int!) {
          distinctPropertyValuesInElementGroupByName(
            elementGroupId: $elementGroupId
            name: "Category"
            filter: { query: "'property.name.Element Context'==Instance" }
          ) {
            results {
              values(limit: $limit) {
                value
                count
              }
            }
          }
        }
        """

        variables = {
            "elementGroupId": group_id,
            "limit": 1000,  # High limit to get all categories
        }

        try:
            data = execute_graphql(query, token, region, variables)
            block = data.get("distinctPropertyValuesInElementGroupByName") or {}
            results_list = block.get("results") or []

            # Create a dictionary of category counts from the model
            model_category_counts = {}
            for r in results_list:
                values = r.get("values") or []
                for v in values:
                    category_name = v.get("value", "")
                    element_count = v.get("count", 0)
                    if category_name:
                        model_category_counts[category_name] = element_count

        except Exception as e:
            raise vkt.UserError(f"Failed to fetch categories from model: {str(e)}")

        vkt.progress_message("Preparing category data summary...", percentage=80)

        # Create main data group
        main_group = vkt.DataGroup()

        # Add summary statistics
        total_categories = len(all_categories)
        categories_in_model = sum(
            1 for cat in all_categories if model_category_counts.get(cat, 0) > 0
        )
        categories_in_contract = len(required_categories)
        categories_matched = sum(
            1 for cat in required_categories if model_category_counts.get(cat, 0) > 0
        )

        summary_group = vkt.DataGroup(
            vkt.DataItem("Total Categories", total_categories),
            vkt.DataItem("Categories in Model", categories_in_model),
            vkt.DataItem("Categories in Contract", categories_in_contract),
            vkt.DataItem(
                "Contract Categories Found",
                categories_matched,
                status=vkt.DataStatus.SUCCESS
                if categories_matched == categories_in_contract
                else vkt.DataStatus.WARNING,
            ),
        )
        main_group.add(vkt.DataItem("Summary", subgroup=summary_group))

        # Add category details grouped by status
        present_group = vkt.DataGroup()
        missing_from_model_group = vkt.DataGroup()
        missing_from_contract_group = vkt.DataGroup()
        not_applicable_group = vkt.DataGroup()

        for category_name in all_categories:
            # Check if category is in the model
            element_count = model_category_counts.get(category_name, 0)
            in_model = element_count > 0

            # Check if category is in required categories
            in_contract = category_name in required_categories

            # Categorize and add to appropriate group
            if in_contract and in_model:
                # Present in both contract and model - SUCCESS
                present_group.add(
                    vkt.DataItem(
                        category_name,
                        element_count,
                        suffix="elements",
                        status=vkt.DataStatus.SUCCESS,
                        status_message="✓ Present in contract and model",
                    )
                )
            elif in_contract and not in_model:
                # In contract but not in model - ERROR
                missing_from_model_group.add(
                    vkt.DataItem(
                        category_name,
                        "0",
                        suffix="elements",
                        status=vkt.DataStatus.ERROR,
                        status_message="✗ In contract but not in model",
                    )
                )
            elif not in_contract and in_model:
                # In model but missing from contract - WARNING
                missing_from_contract_group.add(
                    vkt.DataItem(
                        category_name,
                        element_count,
                        suffix="elements",
                        status=vkt.DataStatus.WARNING,
                        status_message="✗ Missing in the contract",
                    )
                )
            else:
                # Not in contract and not in model - INFO
                not_applicable_group.add(
                    vkt.DataItem(
                        category_name,
                        "0",
                        suffix="elements",
                        status=vkt.DataStatus.INFO,
                        status_message="Not in contract, not in model",
                    )
                )

        # Add grouped categories to main group
        if len(present_group) > 0:
            main_group.add(
                vkt.DataItem("✓ Present (Contract & Model)", subgroup=present_group)
            )

        if len(missing_from_model_group) > 0:
            main_group.add(
                vkt.DataItem("✗ Missing from Model", subgroup=missing_from_model_group)
            )

        if len(missing_from_contract_group) > 0:
            main_group.add(
                vkt.DataItem(
                    "⚠ Missing from Contract", subgroup=missing_from_contract_group
                )
            )

        if len(not_applicable_group) > 0:
            main_group.add(
                vkt.DataItem("○ Not Applicable", subgroup=not_applicable_group)
            )

        return vkt.DataResult(main_group)

    def download_category_report(self, params, **kwargs):
        """
        Generate and download a Word document containing the Category Summary table.

        Args:
            params: User input parameters
            **kwargs: Additional arguments

        Returns:
            DownloadResult with the Word document
        """
        import io
        from datetime import datetime

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor

        if not params.autodesk_file:
            raise vkt.UserError(
                "Please select an Autodesk file from the input field above"
            )

        # Initialize the OAuth2 integration and get access token
        integration = vkt.external.OAuth2Integration("autodesk-integration")
        token = integration.get_access_token()

        # Get region and AEC Data Model element group ID from the Autodesk file
        region = params.autodesk_file.get_region(token)
        group_id = params.autodesk_file.get_aec_data_model_element_group_id(token)

        # Extract required categories from dynamic array
        required_categories = set(row["category"] for row in params.required_categories)

        # Define the master list of categories (same as dropdown options)
        all_categories = [
            "Structural Framing",
            "Structural Columns",
            "Structural Foundations",
            "Walls",
            "Floors",
            "Roofs",
            "Ceilings",
            "Doors",
            "Windows",
            "Stairs",
            "Railings",
            "Curtain Panels",
            "Curtain Wall Mullions",
            "Furniture",
            "Mechanical Equipment",
            "Plumbing Fixtures",
            "Lighting Fixtures",
            "Electrical Equipment",
            "Ducts",
            "Pipes",
        ]

        vkt.progress_message("Fetching category data from model...", percentage=20)

        # Query to get all distinct categories in the model with their counts
        query = """
        query UsedCategories($elementGroupId: ID!, $limit: Int!) {
          distinctPropertyValuesInElementGroupByName(
            elementGroupId: $elementGroupId
            name: "Category"
            filter: { query: "'property.name.Element Context'==Instance" }
          ) {
            results {
              values(limit: $limit) {
                value
                count
              }
            }
          }
        }
        """

        variables = {
            "elementGroupId": group_id,
            "limit": 1000,  # High limit to get all categories
        }

        try:
            data = execute_graphql(query, token, region, variables)
            block = data.get("distinctPropertyValuesInElementGroupByName") or {}
            results_list = block.get("results") or []

            # Create a dictionary of category counts from the model
            model_category_counts = {}
            for r in results_list:
                values = r.get("values") or []
                for v in values:
                    category_name = v.get("value", "")
                    element_count = v.get("count", 0)
                    if category_name:
                        model_category_counts[category_name] = element_count

        except Exception as e:
            raise vkt.UserError(f"Failed to fetch categories from model: {str(e)}")

        vkt.progress_message("Generating Word document...", percentage=60)

        # Create Word document
        doc = Document()

        # Add title
        title = doc.add_heading("Category Summary Report", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(
            f"File: {params.autodesk_file.url.split('/')[-1] if params.autodesk_file else 'Unknown'}"
        )
        doc.add_paragraph("")

        # Add legend
        doc.add_heading("Legend", level=2)
        legend_items = [
            ("✓ Green", "Category is in contract and present in model"),
            ("✗ Orange", "Category is in contract but not in model"),
            ("✗ Red", "Category is in model but missing from contract"),
            ("✗ Gray", "Category is neither in contract nor in model"),
        ]
        for symbol, description in legend_items:
            doc.add_paragraph(f"{symbol}: {description}", style="List Bullet")

        doc.add_paragraph("")

        # Add table
        doc.add_heading("Category Details", level=2)

        # Create table with 3 columns
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Category"
        header_cells[1].text = "Status"
        header_cells[2].text = "Description"

        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Add data rows
        for category_name in all_categories:
            # Check if category is in the model
            element_count = model_category_counts.get(category_name, 0)
            in_model = element_count > 0

            # Check if category is in required categories
            in_contract = category_name in required_categories

            # Determine status symbol, description, and color
            if in_contract and in_model:
                status_symbol = "✓"
                status_text = f"Present ({element_count} elements)"
                color = RGBColor(0, 128, 0)  # Green
            elif in_contract and not in_model:
                status_symbol = "✗"
                status_text = "In contract but not in model"
                color = RGBColor(255, 165, 0)  # Orange
            elif not in_contract and in_model:
                status_symbol = "✗"
                status_text = "Missing in the contract"
                color = RGBColor(255, 0, 0)  # Red
            else:  # not in_contract and not in_model
                status_symbol = "✗"
                status_text = "Not in contract, not in model"
                color = RGBColor(128, 128, 128)  # Gray

            # Add row to table
            row_cells = table.add_row().cells
            row_cells[0].text = category_name
            row_cells[1].text = status_symbol
            row_cells[2].text = status_text

            # Apply color to status symbol
            for paragraph in row_cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = color
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        vkt.progress_message("Finalizing document...", percentage=90)

        # Save document to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Create filename with timestamp
        filename = f"Category_Summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        # Return as DownloadResult
        return vkt.DownloadResult(vkt.File.from_data(doc_io.getvalue()), filename)
